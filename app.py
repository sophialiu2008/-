import streamlit as st
from dashscope import MultiModalConversation, Generation
from dashscope.audio.tts import SpeechSynthesizer
import dashscope
from PIL import Image, ImageDraw, ImageFont
import tempfile
import os
import qrcode
import io
import requests
import docx  # 用于生成Word
from docx.shared import Pt, RGBColor # 用于调整Word字体样式
import PyPDF2
import math

# --- 1. 页面配置 ---
st.set_page_config(
    page_title="小学语文作文批改宝",
    page_icon="🎓",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# --- 🎨 核心美化：自定义 CSS ---
st.markdown("""
    <style>
    .stApp { background-color: #f8f9fa; font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    h1 { color: #2c3e50; font-weight: 700; text-align: center; padding-bottom: 10px; border-bottom: 3px solid #FF9F43; margin-bottom: 20px; font-size: 1.8rem !important; }
    [data-testid="stSidebar"] { background-color: #ffffff; box-shadow: 2px 0 5px rgba(0,0,0,0.05); }
    .stButton>button { width: 100%; border-radius: 25px; height: 50px; font-weight: bold; border: none; background: linear-gradient(135deg, #FF9F43 0%, #ff6b6b 100%); color: white; box-shadow: 0 4px 6px rgba(0,0,0,0.1); transition: all 0.3s ease; }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 6px 8px rgba(0,0,0,0.15); color: white !important; }
    div[data-testid="column"] .stButton>button { background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); }
    .stTextArea textarea { background-color: #fffdf5; border: 1px solid #e0e0e0; border-radius: 10px; font-size: 16px; line-height: 1.6; }
    .stSuccess, .stInfo, .stWarning, .stError { border-radius: 10px; border: none; box-shadow: 0 2px 5px rgba(0,0,0,0.05); }
    .stSpinner > div { border-top-color: #FF9F43 !important; }
    .upload-hint { font-size: 0.9rem; color: #d35400; background-color: #ffe0b2; padding: 12px; border-radius: 12px; margin-bottom: 15px; border-left: 5px solid #e67e22; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
    </style>
    """, unsafe_allow_html=True)

# --- 顶部 Header ---
col_logo, col_title = st.columns([1, 5])
with col_logo: st.markdown("## 🎓")
with col_title: st.markdown("<h1>小学语文作文批改宝</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #7f8c8d; margin-top: -15px;'>🚀 以评促写 | 温暖引导 | 深度思维启发</p>", unsafe_allow_html=True)

# --- 2. 基础配置 ---
api_key = st.secrets.get("DASHSCOPE_API_KEY")
if not api_key:
    st.error("⚠️ 请配置 API Key")
    st.stop()
dashscope.api_key = api_key

if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = ""
if 'review_result' not in st.session_state:
    st.session_state.review_result = ""

# --- 🛠️ 核心工具函数 ---
def compress_image(image, max_width=1024):
    if image.width > max_width:
        ratio = max_width / image.width
        new_height = int(image.height * ratio)
        return image.resize((max_width, new_height), Image.Resampling.LANCZOS)
    return image

# 🌟 更新：删除了童声和新闻播报
def generate_audio_dashscope(text, voice_name):
    voice_map = {
        "👩‍🏫 温柔女老师 (知厨)": "sambert-zhichu-v1",
        "👨‍🏫 阳光男老师 (知达)": "sambert-zhida-v1"
    }
    model_id = voice_map.get(voice_name, "sambert-zhichu-v1")
    try:
        text = text.replace("**", "").replace("###", "").replace("---", "")
        if len(text) > 800: text = text[:800]
        result = SpeechSynthesizer.call(model=model_id, text=text, sample_rate=48000)
        if result.get_audio_data() is not None:
            with open("review.mp3", "wb") as f:
                f.write(result.get_audio_data())
            return True
        return False
    except Exception as e:
        st.warning(f"语音服务繁忙: {e}")
        return False

@st.cache_resource
def get_font():
    font_path = "SimHei.ttf"
    if not os.path.exists(font_path):
        url = "https://github.com/StellarCN/scp_zh/raw/master/fonts/SimHei.ttf"
        try:
            with st.spinner("🚀 正在初始化字体资源..."):
                r = requests.get(url, timeout=10)
                if r.status_code == 200:
                    with open(font_path, "wb") as f:
                        f.write(r.content)
        except: return None 
    return font_path

# 🌟 新增：生成 Word 文档工具
def create_word_report(text):
    doc = docx.Document()
    
    # 标题
    title = doc.add_heading('🏆 小学语文作文批改报告', 0)
    title.alignment = 1 # 居中

    # 解析 Markdown 文本并简单的转为 Word 格式
    lines = text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
            
        if clean_line.startswith('###'):
            # 处理小标题 (例如 ### 🌟 亮点)
            heading_text = clean_line.replace('#', '').strip()
            doc.add_heading(heading_text, level=2)
        elif clean_line.startswith('**') and clean_line.endswith('**'):
            # 处理加粗行
            p = doc.add_paragraph()
            run = p.add_run(clean_line.replace('*', ''))
            run.bold = True
        else:
            # 普通段落
            doc.add_paragraph(clean_line)

    # 底部水印
    p = doc.add_paragraph('\nGenerated by 小学语文作文批改宝 (AI)')
    p.alignment = 2 # 右对齐
    
    # 保存到内存
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def create_review_card(text):
    font_path = get_font()
    try:
        title_font = ImageFont.truetype(font_path, 40) if font_path else ImageFont.load_default()
        content_font = ImageFont.truetype(font_path, 24) if font_path else ImageFont.load_default()
    except:
        title_font = ImageFont.load_default()
        content_font = ImageFont.load_default()

    chars_per_line = 32
    line_height = 35
    margin = 40
    header_height = 120
    footer_height = 80
    
    display_lines = []
    paragraphs = text.split('\n')
    for para in paragraphs:
        clean_line = para.replace('#', '').replace('*', '')
        if not clean_line.strip():
            display_lines.append("")
            continue
        for i in range(0, len(clean_line), chars_per_line):
            display_lines.append(clean_line[i:i+chars_per_line])
    
    total_content_height = len(display_lines) * line_height
    img_height = header_height + total_content_height + footer_height
    img_width = 800

    img = Image.new('RGB', (img_width, img_height), color=(255, 255, 245))
    draw = ImageDraw.Draw(img)

    draw.text((40, 40), "🏆 作文批改报告", fill=(255, 75, 75), font=title_font)
    draw.line((40, 100, 760, 100), fill=(200, 200, 200), width=2)
    
    y_text = header_height
    for line in display_lines:
        draw.text((margin, y_text), line, fill=(50, 50, 50), font=content_font)
        y_text += line_height
        
    draw.text((margin, img_height - 50), "🤖 小学语文作文批改宝", fill=(150, 150, 150), font=content_font)
    return img

def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])
def read_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages: text += page.extract_text() + "\n"
    return text
def stitch_images(image_list):
    if not image_list: return None
    images = [Image.open(x) for x in image_list]
    widths, heights = zip(*(i.size for i in images))
    new_im = Image.new('RGB', (max(widths), sum(heights)), (255, 255, 255))
    y_offset = 0
    for im in images: new_im.paste(im, (0, y_offset)); y_offset += im.size[1]
    return compress_image(new_im)

# --- 3. 侧边栏 ---
with st.sidebar:
    st.markdown("### ⚙️ 批改设置")
    grade = st.select_slider("选择年级", options=["一/二年级", "三/四年级", "五/六年级"], value="三/四年级")
    
    # 🌟 更新：只保留男女老师
    voice_choice = st.selectbox(
        "🔊 朗读声音",
        ["👩‍🏫 温柔女老师 (知厨)", "👨‍🏫 阳光男老师 (知达)"]
    )
    
    st.markdown("---")
    st.markdown("### 📤 上传文件")
    
    st.markdown("""
    <div class="upload-hint">
        📱 <b>操作小贴士：</b><br>
        • <b>传图片</b>：点方式一，直接拍照或选相册。<br>
        • <b>传文档</b>：点方式二，选“浏览”或“文件”。
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("##### 📸 方式一：图片 (拍照/相册)")
    uploaded_imgs = st.file_uploader("img", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, key="img_uploader", label_visibility="collapsed")
    st.markdown("##### 📄 方式二：文档 (Word/PDF)")
    uploaded_docs = st.file_uploader("doc", type=['docx', 'pdf'], accept_multiple_files=True, key="doc_uploader", label_visibility="collapsed")
    
    st.markdown("---")
    app_url = "https://share.streamlit.io"
    qr = qrcode.QRCode(box_size=5, border=2)
    qr.add_data(app_url)
    qr.make(fit=True)
    st.image(qr.make_image(fill='black', back_color='white').get_image(), caption="手机扫码使用")

# --- 4. 主逻辑 ---
final_file = None
file_type = ""
is_multiple_imgs = False
img_list_to_stitch = []

if uploaded_docs:
    final_file = uploaded_docs[0]
    file_type = final_file.name.split('.')[-1].lower()
elif uploaded_imgs:
    if len(uploaded_imgs) > 1:
        is_multiple_imgs = True
        img_list_to_stitch = uploaded_imgs
        file_type = "jpg"
    else:
        final_file = uploaded_imgs[0]
        file_type = final_file.name.split('.')[-1].lower()

if final_file or is_multiple_imgs:
    st.markdown("---")
    # === 分支1：处理图片 ===
    if is_multiple_imgs or file_type in ['png', 'jpg', 'jpeg']:
        if is_multiple_imgs:
            st.info(f"📸 正在自动拼接 {len(uploaded_imgs)} 张图片...")
            image = stitch_images(img_list_to_stitch) 
            file_name_for_tmp = "stitched.jpg"
        else:
            image = Image.open(final_file)
            image = compress_image(image)
            file_name_for_tmp = final_file.name
            
        st.image(image, caption='预览(已自动压缩)', use_container_width=True)
        
        file_suffix = os.path.splitext(file_name_for_tmp)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_suffix) as tmp_file:
            image.save(tmp_file)
            tmp_file_path = tmp_file.name

        if st.button("🔍 开始识别文字"):
            with st.spinner('👀 正在识别手写字迹...'):
                try:
                    msg = [{'role': 'user', 'content': [{'image': f"file://{tmp_file_path}"}, {'text': 'OCR识别，仅输出作文正文，不要输出任何其他内容。'}]}]
                    resp = MultiModalConversation.call(model='qwen-vl-max', messages=msg)
                    if resp.status_code == 200:
                        st.session_state.extracted_text = resp.output.choices[0].message.content[0]['text']
                        st.rerun()
                except Exception as e: st.error(f"错误: {e}")

    # === 分支2：处理文档 ===
    elif file_type in ['docx', 'pdf']:
        if st.button("📖 读取文档内容"):
            try:
                if file_type == 'docx': st.session_state.extracted_text = read_docx(final_file)
                else: st.session_state.extracted_text = read_pdf(final_file)
                st.rerun()
            except Exception as e: st.error(f"读取失败: {e}")

    # === 公共部分：批改与展示 ===
    if st.session_state.extracted_text:
        st.subheader("📝 作文内容确认")
        user_text = st.text_area("内容", value=st.session_state.extracted_text, height=200, label_visibility="collapsed")
        
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("✨ 智能批改 (深度评估版)"):
            with st.spinner('⚡ 老师正在从基础规范、内容表达、思维情感、创意个性四个维度进行评估...'):
                
                grade_focus = ""
                if grade == "一/二年级":
                    grade_focus = "侧重【敢写、能写、写完整】。鼓励为主，不苛求文采，重点关注字词基础和句子是否通顺。"
                elif grade == "三/四年级":
                    grade_focus = "侧重【写清楚、有细节、有顺序】。引导学生把过程写具体，注意段落结构。"
                else: 
                    grade_focus = "侧重【有中心、有情感、有思考】。鼓励个性表达，关注思想深度和文学意识。"

                prompt = f"""
                你是一位秉持“以评促写、以评育人”理念的小学语文老师。请根据以下标准批改这篇{grade}学生的作文。

                **【评价标准】（综合考虑以下四个维度）**
                1. **基础规范 (30%)**：字迹/格式、标点使用、语句通顺（无明显语病）、用词恰当、无错别字。
                2. **内容表达 (30%)**：切题与中心明确、内容具体（拒绝空洞套话）、条理清晰（有逻辑联系）。
                3. **思维与情感 (20%)**：真情实感（拒绝成人化模板）、有观察与思考（能发现细节，有感悟）。
                4. **创意与个性 (20%)**：语言有特色（修辞）、构思新颖、有想象力。

                **【当前年级侧重】**：{grade_focus}

                **【作文内容】**：
                {user_text}

                **【批改要求】**
                1. **拒绝冷冰冰的判词**：评语要具体、温暖、有引导性。例如：“你把...写得特别生动，如果能...就更棒了！”。
                2. **输出格式**：请严格按以下Markdown格式输出：
                   ### 🌟 亮点与光芒
                   (挖掘孩子作文中属于“思维情感”和“创意个性”的闪光点，给予具体表扬)
                   
                   ### 🩺 基础诊疗室
                   (指出“基础规范”方面的问题，如错别字、标点、病句，并给出正确写法)
                   
                   ### 💡 提升小锦囊
                   (针对“内容表达”提出建议，如如何让内容更具体、条理更清晰，给出1-2个具体的修改示范)
                   
                   ### 🏆 综合评价
                   (给出等级：A+/A/B，并写一句温暖的寄语，鼓励孩子继续写作)
                """
                
                try:
                    resp = Generation.call(model='qwen-turbo', messages=[{'role': 'user', 'content': prompt}])
                    if resp.status_code == 200:
                        st.session_state.review_result = resp.output.text
                        st.balloons()
                    else: st.error("失败")
                except Exception as e: st.error(f"错误: {e}")

        if st.session_state.review_result:
            st.markdown("---")
            st.subheader("📝 批改结果")
            with st.container():
                st.markdown(st.session_state.review_result)
            
            st.markdown("---")
            st.subheader("🎁 互动功能")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔊 播放语音点评"):
                    with st.spinner(f"正在生成语音..."):
                        if generate_audio_dashscope(st.session_state.review_result, voice_choice):
                            st.audio("review.mp3")
            with c2:
                # 🌟 核心更新：这里增加了 Word 下载按钮，并将其与图片下载放在一起
                col_word, col_img = st.columns(2)
                
                with col_word:
                    # 生成 Word 文档
                    word_file = create_word_report(st.session_state.review_result)
                    st.download_button(
                        label="📄 下载Word文档",
                        data=word_file,
                        file_name="作文批改报告.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                with col_img:
                    # 生成图片
                    img = create_review_card(st.session_state.review_result)
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    st.download_button(
                        label="🖼️ 下载评语图片", 
                        data=buf.getvalue(), 
                        file_name="评语.png", 
                        mime="image/png"
                    )

else:
    st.info("👈 请点击左上角箭头打开侧边栏，上传作文图片或文档。")
